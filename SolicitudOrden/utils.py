# =============================================================================
# MÓDULO UTILS: Funciones de Procesamiento, Conversión y Comunicación
# =============================================================================

# -------------------------------------------------------------------------
# LIBRERÍAS ESTÁNDAR
# -------------------------------------------------------------------------
import math
import os
import re
import glob
import subprocess
import mimetypes
import smtplib
from collections import Counter

from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

# -------------------------------------------------------------------------
# LIBRERÍAS DE TERCEROS
# -------------------------------------------------------------------------
import openpyxl
from openpyxl.worksheet.worksheet import Worksheet
from openpyxl.cell.cell import MergedCell
from openpyxl.styles import Alignment

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

from PyPDF2 import PdfReader, PdfWriter
import language_tool_python

# -------------------------------------------------------------------------
# MÓDULOS LOCALES / CONFIG
# -------------------------------------------------------------------------
from config import CAMPOS, HOJA_SOLICITUD, OBLIGACIONES

# Inicialización LanguageTool
tool = language_tool_python.LanguageTool('es')

# --------------------------
# UTIL: sanitizar nombres
# --------------------------
def sanitize_filename(name: str) -> str:
    """Quita caracteres no válidos para filenames y colapsa espacios."""
    if name is None:
        return ""
    name = str(name).strip()
    # Reemplaza caracteres problemáticos por guion bajo
    name = re.sub(r'[\\/*?:"<>|]', '_', name)
    # Normaliza espacios múltiples
    name = re.sub(r'\s+', ' ', name)
    return name

# =============================================================================
# 1. CORRECCIÓN Y CARGA
# =============================================================================

def corregir_texto(texto: str) -> str:
    try:
        if texto is None:
            return ""
        texto = str(texto).strip()
        if not texto:
            return ""
        matches = tool.check(texto)
        return language_tool_python.utils.correct(texto, matches)
    except Exception as e:
        print(f"⚠️ language_tool fallo: {e}")
        return texto

def cargar_plantilla(ruta_plantilla):
    return openpyxl.load_workbook(ruta_plantilla)

def cargar_datos(ruta_datos):
    return openpyxl.load_workbook(ruta_datos)

def buscar_y_reemplazar(hoja, texto_buscar, texto_reemplazar):
    if texto_reemplazar is None:
        texto_reemplazar = ""
    texto_reemplazar = str(texto_reemplazar).strip()
    for row in hoja.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            if isinstance(cell.value, str) and texto_buscar in cell.value:
                cell.value = cell.value.replace(texto_buscar, texto_reemplazar)

# =============================================================================
# 2. EXCEL - PROCESAMIENTO / GUARDADO
# =============================================================================

def ajustar_tamano_celdas(hoja: Worksheet, celdas_a_ajustar: list,
                          ancho_max_columna: int = 197, factor_ajuste: float = 1.05):
    filas_procesadas = set()
    for fila, col in celdas_a_ajustar:
        celda = hoja.cell(row=fila, column=col)
        valor_celda = str(celda.value) if celda.value else ""
        if fila in filas_procesadas or not valor_celda:
            continue
        celda.alignment = Alignment(wrapText=True, horizontal='left', vertical='center')
        caracteres_por_linea = int(ancho_max_columna * factor_ajuste)
        if caracteres_por_linea > 0:
            num_lineas = math.ceil(len(valor_celda) / caracteres_por_linea)
            nueva_altura = num_lineas * 14
            nueva_altura = max(15, min(nueva_altura, 300))
            hoja.row_dimensions[fila].height = nueva_altura
            filas_procesadas.add(fila)

def procesar_solicitud(wb_plantilla, datos_fila):
    hoja = wb_plantilla[HOJA_SOLICITUD]
    for campo_db, campo_plantilla in CAMPOS.items():
        if campo_db not in OBLIGACIONES:
            valor = datos_fila.get(campo_db, '')
            valor = corregir_texto(str(valor))
            buscar_y_reemplazar(hoja, campo_plantilla, valor)
    obligaciones_reales = [
        corregir_texto(str(datos_fila.get(campo)).strip())
        for campo in OBLIGACIONES
        if datos_fila.get(campo) and str(datos_fila.get(campo)).strip() != ""
    ]
    num_obligaciones = len(obligaciones_reales)
    if num_obligaciones <= 5:
        for i, obligacion in enumerate(obligaciones_reales, start=1):
            marcador_excel = f"{{obligacion_{i}}}"
            buscar_y_reemplazar(hoja, marcador_excel, obligacion)
        for i in range(num_obligaciones + 1, 6):
            marcador_excel = f"{{obligacion_{i}}}"
            buscar_y_reemplazar(hoja, marcador_excel, "")
    else:
        for i in range(1, 6):
            marcador_excel = f"{{obligacion_{i}}}"
            buscar_y_reemplazar(hoja, marcador_excel, " ")
    return wb_plantilla

def guardar_solicitud(wb, nombre_archivo, ruta_output):
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell, MergedCell):
                    continue
                if cell.value is None:
                    cell.value = ""
                elif isinstance(cell.value, str) and cell.value.strip().lower() == "none":
                    cell.value = ""
    ruta_completa = os.path.join(ruta_output, nombre_archivo)
    # Si ya existe, añade sufijo incremental para evitar sobreescritura accidental
    base, ext = os.path.splitext(ruta_completa)
    contador = 1
    while os.path.exists(ruta_completa):
        ruta_completa = f"{base}_{contador}{ext}"
        contador += 1
    wb.save(ruta_completa)
    return ruta_completa

# =============================================================================
# 3. WORD (Anexo de Obligaciones)
# =============================================================================

def generar_word_obligaciones(datos_fila, ruta_output, nombre_archivo, ruta_plantilla_word):
    obligaciones = []
    for campo in OBLIGACIONES:
        valor = datos_fila.get(campo)
        if valor and str(valor).strip() != "":
            obligaciones.append(corregir_texto(str(valor).strip()))
    if len(obligaciones) <= 5:
        return None
    try:
        doc = Document(ruta_plantilla_word)
    except Exception as e:
        print(f"❌ Error al cargar la plantilla de Word: {e}")
        return None
    for p in doc.paragraphs:
        for campo, valor in datos_fila.items():
            if campo in OBLIGACIONES:
                continue
            marcador_word = f"{{{campo}}}"
            if marcador_word in p.text:
                p.text = p.text.replace(marcador_word, corregir_texto(valor))
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
    for p in doc.paragraphs:
        if "{obligaciones}" in p.text:
            p.clear()
            for i, obligacion in enumerate(obligaciones, start=1):
                new_p = p.insert_paragraph_before(f"{i}. {obligacion}")
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_p.paragraph_format.space_before = Pt(0)
                new_p.paragraph_format.space_after = Pt(0)
                new_p.paragraph_format.first_line_indent = Pt(0)
                for run in new_p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            p._element.getparent().remove(p._element)
            break
    nombre_doc = sanitize_filename(nombre_archivo)
    ruta_doc = os.path.join(ruta_output, nombre_doc)
    # Evitar sobreescritura accidental
    base, ext = os.path.splitext(ruta_doc)
    contador = 1
    while os.path.exists(ruta_doc):
        ruta_doc = f"{base}_{contador}{ext}"
        contador += 1
    doc.save(ruta_doc)
    print(f"Documento Word generado: {ruta_doc}")
    return ruta_doc

# =============================================================================
# 4. CONVERSIÓN A PDF
# =============================================================================

def excel_a_pdf(ruta_excel, ruta_pdf, paginas_to_keep=None, timeout=30):
    if paginas_to_keep is None:
        paginas_to_keep = [0]
    try:
        outdir = os.path.dirname(ruta_pdf) or os.path.dirname(ruta_excel) or "."
        print("DEBUG: Iniciando conversión Excel a PDF vía LibreOffice...")
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf", "--outdir",
            outdir, ruta_excel
        ], check=True, timeout=timeout)
        # LibreOffice genera un pdf con el mismo basename que el xlsx
        expected_name = os.path.splitext(os.path.basename(ruta_excel))[0] + ".pdf"
        generated_pdf = os.path.join(outdir, expected_name)
        if not os.path.exists(generated_pdf):
            # Buscar por coincidencia de 'stem' (case-insensitive) en outdir
            stem = os.path.splitext(os.path.basename(ruta_excel))[0].lower()
            candidates = [p for p in glob.glob(os.path.join(outdir, "*.pdf")) if os.path.splitext(os.path.basename(p))[0].lower() == stem]
            if candidates:
                generated_pdf = candidates[0]
            else:
                print(f"❌ No se encontró el PDF generado por LibreOffice en: {generated_pdf}")
                return
        # Leer y filtrar páginas
        with open(generated_pdf, "rb") as f:
            reader = PdfReader(f)
            total_pages = len(reader.pages)
            writer = PdfWriter()
            for idx in paginas_to_keep:
                if 0 <= idx < total_pages:
                    writer.add_page(reader.pages[idx])
                else:
                    print(f"⚠️ Índice fuera de rango al filtrar: {idx} (total={total_pages})")
            if not writer.pages:
                writer.add_page(reader.pages[0])
            # Guardar en ruta_pdf (evita sobrescribir generated_pdf si es distinto)
            with open(ruta_pdf, "wb") as f_out:
                writer.write(f_out)
        # Si LibreOffice dejó un pdf temporal distinto, intentar eliminarlo
        try:
            if os.path.abspath(generated_pdf) != os.path.abspath(ruta_pdf) and os.path.exists(generated_pdf):
                os.remove(generated_pdf)
        except Exception as e:
            print(f"⚠️ No se pudo borrar el PDF temporal de LibreOffice: {e}")
        print(f"✅ PDF de Solicitud (filtrado) generado: {ruta_pdf}")
    except subprocess.CalledProcessError as cpe:
        print(f"❌ LibreOffice falló al convertir el archivo (Error de Subproceso): {cpe}")
    except subprocess.TimeoutExpired:
        print("❌ La conversión de Excel a PDF superó el tiempo límite (Timeout).")
    except Exception as e:
        print(f"❌ Error genérico en excel_a_pdf: {e}")

def word_a_pdf(ruta_word, ruta_pdf):
    try:
        convert(ruta_word, ruta_pdf)
        print(f"✅ PDF de Anexo generado: {ruta_pdf}")
    except Exception as e:
        print(f"❌ Error al convertir Word a PDF: {e}")

# =============================================================================
# 5. EMAIL (lote)
# =============================================================================

def enviar_email_lote(lote_archivos_y_datos: list, remitente_email: str, remitente_password: str, destinatarios_lista: list, ID_BATCH_UNICO: str):
    ASUNTO_BASE = "Solicitud(es) para aval"
    asunto_final = f"{ASUNTO_BASE} Ref: {ID_BATCH_UNICO}"
    num_ordenes = len(lote_archivos_y_datos)
    print(f"\n--- INICIO ENVÍO SMTP (Lote de {num_ordenes} órdenes) ---")

    lista_solicitudes = ""
    archivos_candidatos = []
    destinatarios_header = ', '.join(destinatarios_lista)

    for i, (_, ruta_excel, ruta_word, datos_fila) in enumerate(lote_archivos_y_datos, start=1):
        radi = datos_fila.get('radi', 'N/A')
        nombre_contratista = datos_fila.get('nombre_contratista', 'Contratista')
        lista_solicitudes += f"{i}. CE - {radi} - 2025 - {nombre_contratista}\n"

        ruta_pdf_excel = ruta_excel.replace(".xlsx", ".pdf")
        if os.path.exists(ruta_pdf_excel) and os.path.getsize(ruta_pdf_excel) > 0:
            archivos_candidatos.append(ruta_pdf_excel)

        if ruta_word:
            posible_pdf = ruta_word.replace(".docx", ".pdf")
            if os.path.exists(posible_pdf) and os.path.getsize(posible_pdf) > 0:
                archivos_candidatos.append(posible_pdf)

    if not archivos_candidatos:
        print("❌ Error: No se encontraron archivos PDF válidos para adjuntar. Cancelando envío.")
        return

    # DEBUG: existencia y tamaños
    print("DEBUG: Archivos candidatos (ruta - existe - tamaño bytes):")
    for p in archivos_candidatos:
        try:
            size = os.path.getsize(p) if os.path.exists(p) else -1
        except Exception:
            size = -1
        print(f" - {p} | exists={os.path.exists(p)} | size={size}")

    # --- Preparar nombres únicos para adjuntar (evita colisiones de basename en el mismo correo)
    basenames = [os.path.basename(p) for p in archivos_candidatos]
    counts = Counter(basenames)
    name_occurrence = {}
    attachments = []  # lista de (ruta_real, nombre_a_mostrar)
    for p in archivos_candidatos:
        base = os.path.basename(p)
        if counts[base] > 1:
            idx = name_occurrence.get(base, 0) + 1
            name_occurrence[base] = idx
            name, ext = os.path.splitext(base)
            unique_name = f"{name}({idx}){ext}"
            # sanitize final name
            unique_name = sanitize_filename(unique_name)
        else:
            unique_name = sanitize_filename(base)
        attachments.append((p, unique_name))

    server = None
    try:
        print("DEBUG: Intentando conexión a SMTP (smtp.gmail.com:587)...")
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        print("DEBUG: Conexión TLS establecida.")
        server.login(remitente_email, remitente_password)
        print("✅ AUTENTICACIÓN EXITOSA.")

        msg = MIMEMultipart()
        msg['From'] = remitente_email
        msg['To'] = destinatarios_header
        msg['Subject'] = asunto_final
        cuerpo = f"""Buen día profesor.

De manera atenta remito la(s) siguiente(s) solicitud(es) para su V.º B.º

{lista_solicitudes}
Gracias de antemano. 

Cordialmente,

-- 
Profesional de apoyo 
Cursos de Extensión de Lenguas Extranjeras
Departamento de Lenguas Extranjeras
Facultad de Ciencias Humanas - Sede Bogotá
Universidad Nacional de Colombia
"""
        msg.attach(MIMEText(cuerpo, 'plain'))

        def adjuntar_archivo(msg, ruta_archivo, attachment_name):
            if not ruta_archivo or not os.path.isfile(ruta_archivo):
                print(f"⚠️ No se adjunta (archivo no válido): {ruta_archivo}")
                return False
            tamaño = os.path.getsize(ruta_archivo)
            if tamaño == 0:
                print(f"⚠️ No se adjunta (archivo 0 bytes): {ruta_archivo}")
                return False
            ctype, encoding = mimetypes.guess_type(ruta_archivo)
            if ctype:
                maintype, subtype = ctype.split('/', 1)
            else:
                maintype, subtype = 'application', 'octet-stream'
            with open(ruta_archivo, "rb") as f:
                payload = f.read()
            # Usamos MIMEApplication para ficheros binarios y añadimos filename codificado (RFC2231)
            part = MIMEApplication(payload, _subtype=subtype)
            # Añadimos filename codificado en utf-8 para evitar problemas con acentos
            part.add_header('Content-Disposition', 'attachment', filename=('utf-8', '', attachment_name))
            msg.attach(part)
            print(f"✅ Archivo adjuntado: {ruta_archivo} as {attachment_name}")
            return True

        attached_count = 0
        for ruta_real, nombre_mostrar in attachments:
            if adjuntar_archivo(msg, ruta_real, nombre_mostrar):
                attached_count += 1

        if attached_count == 0:
            print("❌ Ningún archivo pudo ser adjuntado (tamaño inválido o errores). Cancelando envío.")
            return

        server.sendmail(remitente_email, destinatarios_lista, msg.as_string())
        print(f"✅ Lote enviado con éxito a destinatarios '{destinatarios_header}'. Asunto: '{asunto_final}'")
        print(f"✅ Total de adjuntos físicamente adjuntados: {attached_count}")

    except smtplib.SMTPAuthenticationError as e:
        print(f"❌ ERROR CRÍTICO [SMTP AUTH]: {e}")
    except smtplib.SMTPServerDisconnected as e:
        print(f"❌ ERROR CRÍTICO [SMTP DISCONNECT]: {e}")
    except Exception as e:
        print(f"❌ ERROR GENÉRICO: {e}")
    finally:
        if server:
            try:
                server.quit()
                print("DEBUG: Conexión SMTP cerrada.")
            except Exception as e:
                print(f"⚠️ Error al cerrar conexión SMTP: {e}")

    print("--- FIN ENVÍO SMTP ---")